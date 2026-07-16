import os
import sys
import json
import platform
import subprocess
import shutil
import ctypes
import string
import re
from flask import Flask, render_template, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any

# Load environment variables
load_dotenv()


try:
    from faster_whisper import WhisperModel
    HAS_FASTER_WHISPER = True
except ImportError:
    HAS_FASTER_WHISPER = False

app = Flask(__name__)

# Detect Google Colab environment (checks if /content exists, if Google Drive is mounted, or if Colab env var is set)
IS_COLAB = os.path.exists('/content') or os.path.exists('/content/drive') or 'google.colab' in sys.modules or os.environ.get('IS_COLAB') == 'true'

# Use local uploads folder for all environments (use /tmp on Vercel)
IS_VERCEL = os.environ.get('VERCEL') == '1'
if IS_VERCEL:
    base_dir = '/tmp'
else:
    base_dir = os.path.abspath(os.path.dirname(__file__))

UPLOAD_FOLDER = os.path.join(base_dir, 'uploads')
TRANSCRIPTIONS_FOLDER = os.path.join(base_dir, 'transcriptions')
CONFIG_FILE = os.path.join(base_dir, 'config.json')
TEMPLATE_CATALOG_FILE = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'templates.json')

DEFAULT_AI_MODEL = os.getenv('OPENROUTER_MODEL', 'openai/gpt-5.6-luna')
MODEL_ALIASES = {
    'gemini-2.5-flash': 'google/gemini-2.5-flash',
    'gemini-2.5-pro': 'google/gemini-2.5-pro',
    'gemini-2.0-flash': 'google/gemini-2.0-flash',
    'gpt-5.6-luna': 'openai/gpt-5.6-luna'
}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TRANSCRIPTIONS_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['TRANSCRIPTIONS_FOLDER'] = TRANSCRIPTIONS_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB max request size

# Global error handlers — always return JSON so the frontend never gets HTML error pages
@app.errorhandler(413)
def request_entity_too_large(e):
    return jsonify({'error': 'Request payload is too large. Please reduce the number of sentences or file size.'}), 413

@app.errorhandler(500)
def internal_server_error(e):
    return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.errorhandler(Exception)
def unhandled_exception(e):
    import traceback
    print(f"Unhandled exception:\n{traceback.format_exc()}")
    return jsonify({'error': f'Unexpected server error: {str(e)}'}), 500

# Configuration Helpers
def load_prompt_template(filename):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(base_dir, 'markdown_files', filename)
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def load_template_document():
    """Load global visual guidance and the canonical template catalog."""
    with open(TEMPLATE_CATALOG_FILE, 'r', encoding='utf-8') as f:
        document = json.load(f)

    if not isinstance(document, dict) or not document:
        raise ValueError('templates.json must contain a non-empty JSON object.')

    guidelines = document.get('__guidelines', {})
    if guidelines and not isinstance(guidelines, dict):
        raise ValueError("templates.json '__guidelines' must be a JSON object.")

    catalog = {
        template_id: template
        for template_id, template in document.items()
        if not template_id.startswith('__')
    }
    if not catalog:
        raise ValueError('templates.json must contain at least one template.')

    for template_id, template in catalog.items():
        if not isinstance(template_id, str) or not template_id.strip():
            raise ValueError('Every template must have a non-empty string identifier.')
        if not isinstance(template, dict):
            raise ValueError(f"Template '{template_id}' must be a JSON object.")
        for required_field in ('category', 'density', 'when_to_use', 'constraints'):
            if required_field not in template:
                raise ValueError(f"Template '{template_id}' is missing '{required_field}'.")

    return guidelines, catalog

def load_template_catalog():
    """Load only selectable templates for the AI enum and UI dropdown."""
    _, catalog = load_template_document()
    return catalog

def load_template_guidelines():
    """Load reinforced global and template-specific selection guidance."""
    guidelines, _ = load_template_document()
    return guidelines

def normalize_model_name(model_name):
    model_name = (model_name or DEFAULT_AI_MODEL).strip()
    return MODEL_ALIASES.get(model_name, model_name)

def get_openrouter_api_key(data=None):
    data = data or {}
    return (
        request.headers.get('X-OpenRouter-Key')
        or request.headers.get('X-Gemini-Key')  # Legacy client compatibility
        or data.get('api_key')
        or os.getenv('OPENROUTER_API_KEY')
        or os.getenv('GEMINI_API_KEY')  # Legacy environment compatibility
    )

class OpenRouterError(Exception):
    def __init__(self, message, status_code=500):
        super().__init__(message)
        self.status_code = status_code

def call_openrouter_api(model_name, prompt, response_schema=None, api_key=None):
    import urllib.request
    import urllib.error
    import json
    
    if not api_key:
        api_key = os.getenv('OPENROUTER_API_KEY') or os.getenv('GEMINI_API_KEY')
        
    if not api_key:
        raise OpenRouterError("OpenRouter API Key is missing. Please configure OPENROUTER_API_KEY in your environment/UI.", 400)
        
    url = "https://openrouter.ai/api/v1/chat/completions"
    
    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {api_key}',
        'HTTP-Referer': 'https://github.com/AnshulUpgrad/PPS-Timestamp-AI-Tool-Upgrad-V2',
        'X-Title': 'PPS Timestamp AI Tool'
    }
    
    body = {
        "model": model_name,
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ],
        "max_tokens": 4096
    }
    
    if response_schema:
        body["response_format"] = {
            "type": "json_schema",
            "json_schema": {
                "name": "structured_response",
                "strict": True,
                "schema": response_schema
            }
        }
        
    def send_request(request_body):
        req = urllib.request.Request(
            url,
            data=json.dumps(request_body).encode('utf-8'),
            headers=headers,
            method='POST'
        )
        with urllib.request.urlopen(req, timeout=180) as response:
            return json.loads(response.read().decode('utf-8'))

    try:
        response_data = send_request(body)
            
        choices = response_data.get('choices', [])
        if not choices:
            raise OpenRouterError(f"No response choices received from OpenRouter. Response data: {response_data}", 500)
            
        text_response = choices[0].get('message', {}).get('content', '')
        if not text_response:
            raise OpenRouterError(f"Empty content in OpenRouter response. Response data: {response_data}", 500)
            
        return text_response
    except urllib.error.HTTPError as e:
        error_body = e.read().decode('utf-8')
        try:
            error_json = json.loads(error_body)
            error_msg = error_json.get('error', {}).get('message', str(e))
        except Exception:
            error_msg = error_body or str(e)

        # Some OpenRouter providers do not implement strict JSON Schema even
        # though they can still return valid JSON. Fall back to prompt-enforced
        # JSON instead of aborting the user's entire workflow.
        schema_error = any(
            marker in error_msg.lower()
            for marker in ('response_format', 'json_schema', 'structured output', 'schema')
        )
        if response_schema and e.code in (400, 422) and schema_error:
            fallback_body = dict(body)
            fallback_body.pop('response_format', None)
            fallback_body['messages'] = [{
                'role': 'user',
                'content': prompt + '\n\nReturn only a valid JSON object matching the requested structure. Do not use Markdown fences.'
            }]
            try:
                response_data = send_request(fallback_body)
                choices = response_data.get('choices', [])
                text_response = choices[0].get('message', {}).get('content', '') if choices else ''
                if text_response:
                    return text_response
                raise OpenRouterError('OpenRouter returned an empty response during JSON fallback.', 500)
            except OpenRouterError:
                raise
            except Exception as fallback_error:
                raise OpenRouterError(
                    f'OpenRouter structured-output fallback failed: {str(fallback_error)}',
                    500
                )
        raise OpenRouterError(f"OpenRouter HTTP Error ({e.code}): {error_msg}", e.code)
    except Exception as e:
        raise OpenRouterError(f"OpenRouter API call failed: {str(e)}", 500)


def clean_json_response(text_response):
    if not text_response:
        return {}
    
    text = text_response.strip()
    
    # Remove markdown code blocks if present
    if text.startswith("```"):
        # Remove start tag (e.g. ```json or ```)
        first_nl = text.find("\n")
        if first_nl != -1:
            text = text[first_nl:].strip()
        else:
            text = text[3:].strip()
        # Remove end tag (e.g. ```)
        if text.endswith("```"):
            text = text[:-3].strip()
            
    # Try parsing with strict=True first
    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        # Try parsing with strict=False (allows raw newlines/tabs inside strings)
        try:
            return json.loads(text, strict=False)
        except json.JSONDecodeError as e2:
            print("Failed to parse JSON response from LLM.")
            print(f"Raw response:\n{text_response}")
            print(f"Strict parse error: {e}")
            print(f"Lenient parse error: {e2}")
            
            # Try removing trailing commas
            try:
                import re
                cleaned = re.sub(r',\s*([\]}])', r'\1', text)
                return json.loads(cleaned, strict=False)
            except Exception:
                pass
                
            raise ValueError(f"JSON Parsing Error: {str(e2)}. Raw output: {text[:300]}")


# Pydantic schemas for data validation and constraint enforcement
class VisualsItem(BaseModel):
    value: str
    timestamp: float

class VisualsDetail(BaseModel):
    label: str
    value: str
    timestamp: float
    extra: Optional[str] = None

class VisualsContent(BaseModel):
    title: str
    heading_timestamp: float = 0.0
    items: Optional[List[VisualsItem]] = []
    details: Optional[List[VisualsDetail]] = []

class VisualsData(BaseModel):
    template_name: str
    why_chosen: str
    graphics_required: bool
    content: VisualsContent

class SessionKeypoints(BaseModel):
    heading: str
    subheadings: List[str]
    text_content: str
    visuals: VisualsData

class SessionChunk(BaseModel):
    title: str
    summary: str
    sentence_indices: List[int]

class ChunkingResult(BaseModel):
    sessions: List[SessionChunk]

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r') as f:
                return json.load(f)
        except Exception:
            pass
    return {'ffmpeg_path': ''}

def save_config(config):
    try:
        with open(CONFIG_FILE, 'w') as f:
            json.dump(config, f)
        return True
    except Exception:
        return False

def get_ffmpeg_command(custom_path=None):
    if custom_path and os.path.exists(custom_path) and os.path.isfile(custom_path):
        return custom_path
    
    # Try global path
    global_ffmpeg = shutil.which('ffmpeg')
    if global_ffmpeg:
        return global_ffmpeg
        
    return None

# Native Tkinter Dialog Subprocess Helper
def run_native_dialog(script):
    try:
        result = subprocess.run([sys.executable, '-c', script], capture_output=True, text=True, timeout=120)
        return result.stdout.strip()
    except Exception as e:
        print(f"Error opening native dialog: {e}")
        return ""

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/env', methods=['GET'])
def get_env():
    is_render = os.environ.get('RENDER') == 'true' or os.environ.get('RENDER') == '1' or platform.system() != 'Windows'
    return jsonify({
        'is_colab': IS_COLAB,
        'is_vercel': IS_VERCEL,
        'is_render': is_render
    }), 200

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in the request'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file:
        filename = secure_filename(file.filename)
        # Ensure filename is unique to avoid overwriting
        base_name, ext = os.path.splitext(filename)
        counter = 1
        unique_name = filename
        while os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], unique_name)):
            unique_name = f"{base_name}_{counter}{ext}"
            counter += 1
            
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_name)
        file.save(filepath)
        
        size_bytes = os.path.getsize(filepath)
        size_mb = size_bytes / (1024 * 1024)
        
        return jsonify({
            'path': filepath,
            'name': unique_name,
            'size': f"{round(size_mb, 2)} MB",
            'size_bytes': size_bytes
        }), 200
    return jsonify({'error': 'Upload failed'}), 500

@app.route('/api/select-file', methods=['POST'])
def select_file():
    if IS_COLAB or os.environ.get('RENDER') == 'true' or platform.system() != 'Windows':
        return jsonify({
            'mode': 'colab'
        }), 200

    # Local fallback
    script = """
import tkinter as tk
from tkinter import filedialog
root = tk.Tk()
root.withdraw()
root.attributes('-topmost', True)
file_path = filedialog.askopenfilename(
    title="Select Media File",
    filetypes=[
        ("Media Files", "*.mp4 *.webm *.mov *.avi *.mkv *.m4v *.mp3 *.wav *.m4a"),
        ("Video Files", "*.mp4 *.webm *.mov *.avi *.mkv *.m4v"),
        ("Audio Files", "*.mp3 *.wav *.m4a"),
        ("All Files", "*.*")
    ]
)
print(file_path, end="")
root.destroy()
"""
    selected_path = run_native_dialog(script)
    if selected_path and os.path.exists(selected_path) and os.path.isfile(selected_path):
        name = os.path.basename(selected_path)
        size_bytes = os.path.getsize(selected_path)
        size_mb = size_bytes / (1024 * 1024)
        return jsonify({
            'mode': 'local',
            'path': selected_path,
            'name': name,
            'size': f"{round(size_mb, 2)} MB",
            'size_bytes': size_bytes
        }), 200
    return jsonify({'path': '', 'mode': 'local'}), 200

@app.route('/api/settings', methods=['GET', 'POST'])
def handle_settings():
    config = load_config()
    ffmpeg_detected_path = get_ffmpeg_command(config.get('ffmpeg_path'))
    
    if request.method == 'POST':
        data = request.json or {}
        new_path = data.get('ffmpeg_path', '').strip()
        config['ffmpeg_path'] = new_path
        if save_config(config):
            updated_detected = get_ffmpeg_command(new_path)
            return jsonify({
                'message': 'Settings saved successfully',
                'ffmpeg_path': new_path,
                'detected_path': updated_detected,
                'is_valid': updated_detected is not None
            }), 200
        else:
            return jsonify({'error': 'Failed to save config file'}), 500
            
    # GET method
    return jsonify({
        'ffmpeg_path': config.get('ffmpeg_path', ''),
        'detected_path': ffmpeg_detected_path,
        'is_valid': ffmpeg_detected_path is not None
    }), 200

@app.route('/api/config', methods=['GET'])
def get_app_config():
    return jsonify({
        'modal_transcribe_url': os.getenv('MODAL_TRANSCRIBE_URL', ''),
        'default_ai_model': DEFAULT_AI_MODEL
    }), 200

@app.route('/api/templates', methods=['GET'])
def get_template_catalog():
    try:
        guidelines, catalog = load_template_document()
        return jsonify({
            'templates': catalog,
            'template_ids': list(catalog.keys()),
            'guidelines': guidelines,
            'default_ai_model': DEFAULT_AI_MODEL
        }), 200
    except Exception as e:
        return jsonify({'error': f'Failed to load template catalog: {str(e)}'}), 500

@app.route('/api/extract', methods=['POST'])
def extract_audio():
    data = request.json or {}
    video_path = data.get('video_path', '')
    mode = data.get('mode', 'copy') # 'copy' or 'mp3'
    
    config = load_config()
    ffmpeg_cmd = get_ffmpeg_command(config.get('ffmpeg_path'))
    
    if not ffmpeg_cmd:
        return jsonify({
            'error': 'FFmpeg executable not found. Please install FFmpeg or set its custom path in Settings.'
        }), 400
        
    if not video_path:
        return jsonify({'error': 'Missing video_path parameter.'}), 400
        
    if not os.path.exists(video_path) or not os.path.isfile(video_path):
        return jsonify({'error': f"Video file not found at path: {video_path}"}), 400

    video_name = os.path.basename(video_path)
    base_name, ext_raw = os.path.splitext(video_name)
    ext = ext_raw.lower().lstrip('.')
    
    unique_base = secure_filename(base_name)
    if not unique_base:
        unique_base = "extracted_audio"
        
    counter = 1
    
    # 1. STREAM COPY EXTRACTION MODE
    if mode == 'copy':
        if ext in ['webm', 'mkv']:
            out_ext = 'webm'
        elif ext in ['mp3', 'wav', 'm4a']:
            out_ext = ext
        else:
            out_ext = 'm4a'
            
        out_filename = f"{unique_base}_extracted.{out_ext}"
        while os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], out_filename)):
            out_filename = f"{unique_base}_extracted_{counter}.{out_ext}"
            counter += 1
            
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], out_filename)
        cmd = [ffmpeg_cmd, '-y', '-i', video_path, '-vn', '-c:a', 'copy', output_path]
        
        try:
            # Run command natively
            result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            size_mb = os.path.getsize(output_path) / (1024 * 1024)
            return jsonify({
                'message': 'Audio successfully extracted (Stream Copy)',
                'filename': out_filename,
                'original_size_bytes': os.path.getsize(video_path),
                'audio_size_bytes': os.path.getsize(output_path),
                'size_mb': round(size_mb, 2),
                'mode': 'Stream Copy',
                'logs': result.stderr
            }), 200
        except subprocess.CalledProcessError as e:
            # Fall back to MP3 transcoding if stream copy fails (e.g. invalid container configurations)
            print(f"Native stream copy failed. Error: {e.stderr}. Trying MP3 transcode fallback...")
            
            out_filename = f"{unique_base}_extracted.mp3"
            while os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], out_filename)):
                out_filename = f"{unique_base}_extracted_{counter}.mp3"
                counter += 1
                
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], out_filename)
            cmd_fallback = [ffmpeg_cmd, '-y', '-i', video_path, '-vn', '-c:a', 'libmp3lame', '-q:a', '4', output_path]
            
            try:
                result_fb = subprocess.run(cmd_fallback, capture_output=True, text=True, check=True)
                size_mb = os.path.getsize(output_path) / (1024 * 1024)
                return jsonify({
                    'message': 'Audio successfully extracted (MP3 Transcode Fallback)',
                    'filename': out_filename,
                    'original_size_bytes': os.path.getsize(video_path),
                    'audio_size_bytes': os.path.getsize(output_path),
                    'size_mb': round(size_mb, 2),
                    'mode': 'Transcode Fallback (MP3)',
                    'logs': f"Stream copy failed:\n{e.stderr}\n\nRunning Transcode:\n{result_fb.stderr}"
                }), 200
            except subprocess.CalledProcessError as e_transcode:
                return jsonify({
                    'error': 'Native extraction failed completely.',
                    'logs': f"Stream copy failed:\n{e.stderr}\n\nTranscode failed:\n{e_transcode.stderr}"
                }), 500
                
    # 2. EXPLICIT TRANSCODE TO MP3
    else:
        out_filename = f"{unique_base}_extracted.mp3"
        while os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], out_filename)):
            out_filename = f"{unique_base}_extracted_{counter}.mp3"
            counter += 1
            
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], out_filename)
        cmd = [ffmpeg_cmd, '-y', '-i', video_path, '-vn', '-c:a', 'libmp3lame', '-q:a', '4', output_path]
        
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            size_mb = os.path.getsize(output_path) / (1024 * 1024)
            return jsonify({
                'message': 'Audio successfully transcoded to MP3',
                'filename': out_filename,
                'original_size_bytes': os.path.getsize(video_path),
                'audio_size_bytes': os.path.getsize(output_path),
                'size_mb': round(size_mb, 2),
                'mode': 'MP3 Transcode',
                'logs': result.stderr
            }), 200
        except subprocess.CalledProcessError as e:
            return jsonify({
                'error': 'MP3 Transcoding failed.',
                'logs': e.stderr
            }), 500

@app.route('/api/files', methods=['GET'])
def list_files():
    try:
        files = []
        for filename in os.listdir(app.config['UPLOAD_FOLDER']):
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.isfile(file_path):
                if filename.endswith('.json'):
                    continue
                size_mb = os.path.getsize(file_path) / (1024 * 1024)
                transcript_name = f"{filename}.json"
                transcript_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], transcript_name)
                has_transcript = os.path.exists(transcript_path)
                files.append({
                    'name': filename,
                    'size': f"{round(size_mb, 2)} MB",
                    'url': f"/uploads/{filename}",
                    'has_transcript': has_transcript
                })
        # Sort files by modification time: latest first
        files.sort(key=lambda x: os.path.getmtime(os.path.join(app.config['UPLOAD_FOLDER'], x['name'])), reverse=True)
        return jsonify(files), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/uploads/<path:filename>')
def download_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

def get_whisper_model(model_size):
    try:
        print(f"Loading WhisperModel '{model_size}' on CUDA...")
        return WhisperModel(model_size, device="cuda", compute_type="float16")
    except Exception as e:
        print(f"CUDA initialization failed or not supported: {e}. Falling back to CPU with int8.")
        return WhisperModel(model_size, device="cpu", compute_type="int8")

@app.route('/api/transcribe', methods=['POST'])
def transcribe_audio():
    modal_url = os.getenv('MODAL_TRANSCRIBE_URL')
    
    data = request.json or {}
    filename = data.get('filename', '')
    model_size = data.get('model_size', 'base')
    
    modal_token_id = data.get('modal_token_id', '').strip()
    modal_token_secret = data.get('modal_token_secret', '').strip()
    
    if modal_token_id:
        os.environ['MODAL_TOKEN_ID'] = modal_token_id
    if modal_token_secret:
        os.environ['MODAL_TOKEN_SECRET'] = modal_token_secret
        
    if not filename:
        return jsonify({'error': 'Missing filename parameter.'}), 400
        
    audio_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(audio_path) or not os.path.isfile(audio_path):
        return jsonify({'error': f"Audio file not found: {filename}"}), 400
        
    transcript_filename = f"{filename}.json"
    transcript_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], transcript_filename)
    
    # Check if we should forward the request to Modal
    use_modal = os.getenv('USE_MODAL', 'true').lower() == 'true' or modal_url
    if use_modal:
        try:
            print(f"Routing transcription for {filename} to Modal via Python SDK...")
            import modal
            
            with open(audio_path, "rb") as f:
                file_bytes = f.read()
                
            print("Looking up WhisperTranscriber Cls...")
            # Use the official Class from_name lookup in Modal
            WhisperTranscriber = modal.Cls.from_name("whisper-transcribe", "WhisperTranscriber")
            server = WhisperTranscriber()
            
            print(f"Calling remote WhisperTranscriber.transcribe method with model_size={model_size}...")
            transcript_data = server.transcribe.remote(audio_bytes=file_bytes, model_name=model_size, language="auto")
            
            with open(transcript_path, 'w', encoding='utf-8') as f:
                json.dump(transcript_data, f, ensure_ascii=False, indent=2)
                
            return jsonify({
                'message': 'Transcription completed successfully via Modal SDK',
                'transcript': transcript_data,
                'transcript_file': transcript_filename
            }), 200
        except Exception as e:
            print(f"Modal native SDK transcription failed: {e}")
            
            # Fall back to HTTP post if they still have the FastAPI endpoint deployed (using their modal_url)
            if modal_url:
                try:
                    # Append query parameters for model and language to modal_url
                    target_url = modal_url
                    delimiter = '&' if '?' in target_url else '?'
                    target_url = f"{target_url}{delimiter}model={model_size}&language=auto"
                    
                    print(f"Falling back to routing to Modal HTTP endpoint: {target_url}")
                    import uuid
                    import urllib.request
                    
                    boundary = f"---Boundary-{uuid.uuid4().hex}"
                    
                    with open(audio_path, "rb") as f:
                        file_bytes = f.read()
                        
                    parts = [
                        f"--{boundary}".encode("utf-8"),
                        f'Content-Disposition: form-data; name="file"; filename="{filename}"'.encode("utf-8"),
                        b"Content-Type: application/octet-stream",
                        b"",
                        file_bytes,
                        f"--{boundary}--".encode("utf-8")
                    ]
                    
                    body = b"\r\n".join(parts)
                    
                    headers = {
                        "Content-Type": f"multipart/form-data; boundary={boundary}",
                        "Content-Length": str(len(body))
                    }
                    
                    req = urllib.request.Request(target_url, data=body, headers=headers, method="POST")
                    with urllib.request.urlopen(req, timeout=600) as response:
                        transcript_data = json.loads(response.read().decode("utf-8"))
                    
                    with open(transcript_path, 'w', encoding='utf-8') as f:
                        json.dump(transcript_data, f, ensure_ascii=False, indent=2)
                        
                    return jsonify({
                        'message': 'Transcription completed successfully via Modal HTTP fallback',
                        'transcript': transcript_data,
                        'transcript_file': transcript_filename
                    }), 200
                except Exception as http_e:
                    print(f"Modal HTTP fallback routing failed: {http_e}")
                    
            if not HAS_FASTER_WHISPER:
                return jsonify({
                    'error': f'Modal transcription failed and local fallback is not available. SDK Error: {str(e)}'
                }), 500
            print("Falling back to local transcription...")

    # Local fallback
    if not HAS_FASTER_WHISPER:
        return jsonify({'error': 'faster-whisper is not installed or available on this server.'}), 500
        
    try:
        model = get_whisper_model(model_size)
        segments, info = model.transcribe(audio_path, beam_size=5, word_timestamps=True)
        
        segments_list = []
        for segment in segments:
            words_list = []
            if segment.words:
                for word in segment.words:
                    words_list.append({
                        'word': word.word,
                        'start': round(word.start, 2),
                        'end': round(word.end, 2),
                        'probability': round(word.probability, 2)
                    })
            segments_list.append({
                'id': segment.id,
                'start': round(segment.start, 2),
                'end': round(segment.end, 2),
                'text': segment.text,
                'words': words_list
            })
            
        transcript_data = {
            'text': ''.join([s['text'] for s in segments_list]).strip(),
            'language': info.language,
            'language_probability': round(info.language_probability, 2),
            'duration': round(info.duration, 2),
            'segments': segments_list
        }
        
        with open(transcript_path, 'w', encoding='utf-8') as f:
            json.dump(transcript_data, f, ensure_ascii=False, indent=2)
            
        return jsonify({
            'message': 'Transcription completed successfully',
            'transcript': transcript_data,
            'transcript_file': transcript_filename
        }), 200
        
    except Exception as e:
        import traceback
        print(f"Whisper transcription failed: {e}\n{traceback.format_exc()}")
        return jsonify({
            'error': 'Whisper transcription failed.',
            'details': str(e)
        }), 500

@app.route('/api/save-transcript', methods=['POST'])
def save_transcript():
    data = request.json or {}
    filename = data.get('filename', '')
    transcript_data = data.get('transcript', {})
    
    if not filename or not transcript_data:
        return jsonify({'error': 'Missing filename or transcript data.'}), 400
        
    try:
        # Create a dummy placeholder in uploads folder so it lists in /api/files
        dummy_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        with open(dummy_path, 'w') as f:
            f.write("")
            
        # Save the transcript JSON
        transcript_filename = f"{filename}.json"
        transcript_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], transcript_filename)
        with open(transcript_path, 'w', encoding='utf-8') as f:
            json.dump(transcript_data, f, ensure_ascii=False, indent=2)
            
        return jsonify({
            'message': 'Transcript saved successfully',
            'filename': filename,
            'transcript_file': transcript_filename
        }), 200
    except Exception as e:
        return jsonify({'error': f"Failed to save transcript: {str(e)}"}), 500

@app.route('/api/transcript/<path:filename>', methods=['GET'])
def get_transcript(filename):
    # Retrieve transcript if it exists
    transcript_filename = f"{filename}.json"
    transcript_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], transcript_filename)
    
    if os.path.exists(transcript_path) and os.path.isfile(transcript_path):
        try:
            with open(transcript_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return jsonify({
                'exists': True,
                'transcript': data
            }), 200
        except Exception as e:
            return jsonify({
                'exists': False,
                'error': f"Failed to load transcript JSON: {str(e)}"
            }), 500
            
    return jsonify({'exists': False}), 200

def split_transcript_into_sentences(transcript_data):
    segments = transcript_data.get('segments', [])
    all_words = []
    for seg in segments:
        for w in seg.get('words', []):
            all_words.append(w)
            
    if not all_words:
        # Fallback to segment-level text if no word-level timestamps exist
        sentences = []
        sentence_id = 0
        for seg in segments:
            text = seg.get('text', '').strip()
            import re
            parts = re.split(r'(?<=[.!?])\s+', text)
            for part in parts:
                if part.strip():
                    sentences.append({
                        'id': sentence_id,
                        'text': part.strip(),
                        'start': seg.get('start', 0.0),
                        'end': seg.get('end', 0.0),
                        'words': []
                    })
                    sentence_id += 1
        return sentences

    sentences = []
    current_words = []
    sentence_id = 0
    
    # Common abbreviations that don't end a sentence
    ABBREVIATIONS = {
        'mr.', 'mrs.', 'ms.', 'dr.', 'prof.', 'sr.', 'jr.', 'vs.', 'etc.', 
        'eg.', 'ie.', 'a.m.', 'p.m.', 'u.s.', 'u.k.', 'jan.', 'feb.', 
        'mar.', 'apr.', 'jun.', 'jul.', 'aug.', 'sep.', 'oct.', 'nov.', 'dec.'
    }
    
    for w in all_words:
        current_words.append(w)
        word_text = w.get('word', '').strip()
        
        # Determine if this word ends a sentence
        is_sentence_end = False
        
        # Check for sentence-ending punctuation
        if word_text and word_text[-1] in ['.', '?', '!']:
            # Check if it's an abbreviation
            clean_word = word_text.lower()
            if clean_word not in ABBREVIATIONS:
                is_sentence_end = True
            
        if is_sentence_end:
            text_parts = []
            for cw in current_words:
                text_parts.append(cw.get('word', ''))
            sentence_text = "".join(text_parts).strip()
            
            sentences.append({
                'id': sentence_id,
                'text': sentence_text,
                'start': current_words[0].get('start', 0.0),
                'end': current_words[-1].get('end', 0.0),
                'words': [{'word': cw.get('word', '').strip(), 'start': cw.get('start', 0.0), 'end': cw.get('end', 0.0)} for cw in current_words]
            })
            sentence_id += 1
            current_words = []
            
    # Capture any trailing words
    if current_words:
        text_parts = []
        for cw in current_words:
            text_parts.append(cw.get('word', ''))
        sentence_text = "".join(text_parts).strip()
        sentences.append({
            'id': sentence_id,
            'text': sentence_text,
            'start': current_words[0].get('start', 0.0),
            'end': current_words[-1].get('end', 0.0),
            'words': [{'word': cw.get('word', '').strip(), 'start': cw.get('start', 0.0), 'end': cw.get('end', 0.0)} for cw in current_words]
        })
        
    return sentences

@app.route('/chunking')
def chunking_view():
    return render_template('chunking.html')

@app.route('/keypoints')
def keypoints_view():
    return render_template('keypoints.html')

def _clamp_timestamp(value, session_start, session_end):
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        numeric = session_start
    return max(session_start, min(session_end, numeric))

def _snap_to_word_start(value, word_starts, session_start, session_end):
    clamped = _clamp_timestamp(value, session_start, session_end)
    if not word_starts:
        return round(clamped, 3)
    return round(min(word_starts, key=lambda start: abs(start - clamped)), 3)

def normalize_generated_visuals(visuals, sentences, session_start, session_end):
    """Keep generated visual timestamps valid, grounded to words, and ordered."""
    word_starts = sorted({
        float(word.get('start', session_start))
        for sentence in sentences
        for word in sentence.get('words', [])
        if session_start <= float(word.get('start', session_start)) <= session_end
    })
    content = visuals.setdefault('content', {})
    heading_time = _snap_to_word_start(
        content.get('heading_timestamp', session_start),
        word_starts,
        session_start,
        session_end
    )
    content['heading_timestamp'] = heading_time

    items = content.get('items') or []
    details = content.get('details') or []
    for entry in items + details:
        entry['timestamp'] = _snap_to_word_start(
            entry.get('timestamp', heading_time),
            word_starts,
            session_start,
            session_end
        )

    items.sort(key=lambda entry: entry.get('timestamp', session_start))
    details.sort(key=lambda entry: entry.get('timestamp', session_start))

    all_entries = sorted(items + details, key=lambda entry: entry.get('timestamp', session_start))
    if all_entries and heading_time >= session_end:
        earlier_word_starts = [start for start in word_starts if start < session_end]
        heading_time = earlier_word_starts[-1] if earlier_word_starts else max(session_start, session_end - 0.1)
        content['heading_timestamp'] = round(heading_time, 3)
    if all_entries and all_entries[0].get('timestamp', heading_time) <= heading_time:
        later_word_starts = [start for start in word_starts if start > heading_time]
        replacement = later_word_starts[0] if later_word_starts else min(session_end, heading_time + 0.1)
        for entry in all_entries:
            if entry.get('timestamp', heading_time) <= heading_time:
                entry['timestamp'] = round(replacement, 3)
        items.sort(key=lambda entry: entry.get('timestamp', session_start))
        details.sort(key=lambda entry: entry.get('timestamp', session_start))

    if str(visuals.get('template_name', '')).startswith('Differentiation Template'):
        # Comparisons remain easy-to-edit paragraph rows: LHS1, RHS1, LHS2, RHS2...
        for index, detail in enumerate(details):
            pair_number = index // 2 + 1
            detail['label'] = f"{'LHS' if index % 2 == 0 else 'RHS'}{pair_number}"

    content['items'] = items
    content['details'] = details
    return visuals

def validate_visual_grounding(session_text, visuals, text_content=''):
    """Reject visual copy that substantially introduces vocabulary absent from the chunk."""
    stop_words = {
        'a', 'an', 'and', 'are', 'as', 'at', 'be', 'by', 'for', 'from', 'in',
        'is', 'it', 'of', 'on', 'or', 'that', 'the', 'this', 'to', 'was', 'with'
    }
    source_tokens = {
        token for token in re.findall(r"[a-z0-9']+", session_text.lower())
        if token not in stop_words
    }
    content = visuals.get('content', {})
    generated_values = [content.get('title', ''), text_content]
    generated_values.extend(item.get('value', '') for item in content.get('items', []))
    generated_values.extend(detail.get('value', '') for detail in content.get('details', []))
    for value in generated_values:
        tokens = [
            token for token in re.findall(r"[a-z0-9']+", str(value).lower())
            if token not in stop_words
        ]
        if len(tokens) < 3:
            continue
        unsupported = [token for token in tokens if token not in source_tokens]
        if len(unsupported) / len(tokens) > 0.34:
            raise ValueError(
                f"Visual text is not sufficiently grounded in the chunk; unsupported terms: {', '.join(unsupported[:8])}"
            )

@app.route('/api/generate-session-keypoints', methods=['POST'])
def generate_session_keypoints():
    data = request.json or {}
    sentences = data.get('sentences', [])
    feedback = data.get('feedback', '').strip()
    existing_heading = data.get('existing_heading', '')
    existing_subheadings = data.get('existing_subheadings', [])
    existing_text_content = data.get('existing_text_content', '')
    existing_visuals = data.get('existing_visuals', None)
    model_name = normalize_model_name(data.get('model'))
        
    # API key from headers, request body, or environment
    api_key = get_openrouter_api_key(data)
    
    if not api_key:
        return jsonify({'error': 'OpenRouter API Key is missing. Please configure OPENROUTER_API_KEY in your .env file.'}), 400
        
    if not sentences:
        return jsonify({'error': 'No sentences provided for generating keypoints.'}), 400
        
    # Construct the session text and timestamped timeline transcript representation
    session_text = " ".join([s.get('text', '') for s in sentences])
    
    timestamped_lines = []
    for s in sentences:
        start_time = s.get('start', 0.0)
        end_time = s.get('end', 0.0)
        words = s.get('words', [])
        if words:
            word_str_list = []
            for w in words:
                w_text = w.get('word', '').strip()
                w_start = w.get('start', 0.0)
                word_str_list.append(f"{w_text}({w_start:.1f}s)")
            words_str = " ".join(word_str_list)
        else:
            words_str = s.get('text', '')
        timestamped_lines.append(f"[{start_time:.1f}s - {end_time:.1f}s]: {words_str}")
        
    timestamped_transcript = "\n".join(timestamped_lines)
    
    session_start = sentences[0].get('start', 0.0) if sentences else 0.0
    session_end = sentences[-1].get('end', 0.0) if sentences else 0.0
    
    # templates.json is the canonical template-selection guide.
    try:
        template_guidelines, template_catalog = load_template_document()
        visuals_guide_content = json.dumps({
            'global_and_reinforced_guidelines': template_guidelines,
            'templates': template_catalog
        }, ensure_ascii=False, indent=2)
    except Exception as e:
        return jsonify({'error': f'Failed to load templates.json: {str(e)}'}), 500
        
    # Design prompt
    if feedback:
        existing_data = {
            "heading": existing_heading,
            "subheadings": existing_subheadings,
            "text_content": existing_text_content
        }
        if existing_visuals:
            existing_data["visuals"] = existing_visuals
            
        existing_str = json.dumps(existing_data, indent=2)
        
        prompt_tmpl = load_prompt_template('keypoints_refinement.md')
        prompt = prompt_tmpl.format(
            feedback=feedback,
            session_text=session_text,
            timestamped_transcript=timestamped_transcript,
            existing_str=existing_str,
            session_start=session_start,
            session_end=session_end,
            visuals_guide_content=visuals_guide_content
        )
    else:
        prompt_tmpl = load_prompt_template('keypoints_initial.md')
        prompt = prompt_tmpl.format(
            session_text=session_text,
            timestamped_transcript=timestamped_transcript,
            session_start=session_start,
            session_end=session_end,
            visuals_guide_content=visuals_guide_content
        )

    schema = {
        "type": "object",
        "properties": {
            "heading": {"type": "string"},
            "subheadings": {
                "type": "array",
                "items": {"type": "string"}
            },
            "text_content": {"type": "string"},
            "visuals": {
                "type": "object",
                "properties": {
                    "template_name": {
                        "type": "string",
                        "enum": list(template_catalog.keys())
                    },
                    "why_chosen": {"type": "string"},
                    "graphics_required": {"type": "boolean"},
                    "content": {
                        "type": "object",
                        "properties": {
                            "title": {"type": "string"},
                            "heading_timestamp": {
                                "type": "number",
                                "minimum": session_start,
                                "maximum": session_end
                            },
                            "items": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "value": {"type": "string"},
                                        "timestamp": {"type": "number"}
                                    },
                                    "required": ["value", "timestamp"],
                                    "additionalProperties": False
                                }
                            },
                            "details": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "label": {"type": "string"},
                                        "value": {"type": "string"},
                                        "timestamp": {"type": "number"}
                                    },
                                    "required": ["label", "value", "timestamp"],
                                    "additionalProperties": False
                                }
                            }
                        },
                        "required": ["title", "heading_timestamp", "items", "details"],
                        "additionalProperties": False
                    }
                },
                "required": ["template_name", "why_chosen", "graphics_required", "content"],
                "additionalProperties": False
            }
        },
        "required": ["heading", "subheadings", "text_content", "visuals"],
        "additionalProperties": False
    }
    
    max_retries = 3
    last_error = None
    
    for attempt in range(max_retries):
        try:
            attempt_prompt = prompt
            if last_error is not None:
                attempt_prompt += (
                    "\n\nYour previous response failed validation: "
                    f"{last_error}. Correct that issue. Use only transcript-supported wording and timestamps."
                )
            text_response = call_openrouter_api(
                model_name=model_name,
                prompt=attempt_prompt,
                response_schema=schema,
                api_key=api_key
            )
            
            result = clean_json_response(text_response)
            
            # Validate response with Pydantic schema
            validated_data = SessionKeypoints.model_validate(result)
            
            heading = validated_data.heading
            subheadings = validated_data.subheadings
            text_content = validated_data.text_content
            visuals = validated_data.visuals.model_dump()
            visuals = normalize_generated_visuals(
                visuals,
                sentences,
                session_start,
                session_end
            )
            template_name = visuals.get('template_name', 'Face Only').strip()
            if template_name == 'Face Only':
                subheadings = []
                text_content = ''
                visuals['graphics_required'] = False
                visuals['content'] = {
                    'title': '',
                    'heading_timestamp': session_start,
                    'items': [],
                    'details': []
                }
            else:
                validate_visual_grounding(session_text, visuals, text_content)
                
            return jsonify({
                'heading': heading,
                'subheadings': subheadings,
                'text_content': text_content,
                'visuals': visuals
            }), 200
        except Exception as e:
            print(f"Keypoints generation attempt {attempt + 1} failed: {e}")
            last_error = e
            import time
            time.sleep(1)
            
    return jsonify({
        'error': f"Failed to run keypoints generation after {max_retries} attempts. Last error: {str(last_error)}"
    }), 500


@app.route('/api/sentences/<path:filename>', methods=['GET'])
def get_sentences(filename):
    transcript_filename = f"{filename}.json"
    transcript_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], transcript_filename)
    
    if not os.path.exists(transcript_path) or not os.path.isfile(transcript_path):
        return jsonify({'error': f"Transcript file not found for: {filename}"}), 404
        
    try:
        with open(transcript_path, 'r', encoding='utf-8') as f:
            transcript_data = json.load(f)
            
        sentences = split_transcript_into_sentences(transcript_data)
        return jsonify({
            'filename': filename,
            'duration': transcript_data.get('duration', 0.0),
            'sentences': sentences
        }), 200
    except Exception as e:
        return jsonify({'error': f"Failed to parse sentences: {str(e)}"}), 500

def reconcile_chunks(sessions, sentences):
    if not sessions or not sentences:
        if sentences:
            # Fallback: create a single chunk containing all sentences
            return [{
                'title': 'Session 1',
                'summary': 'Discussion.',
                'sentence_indices': [s['id'] for s in sentences]
            }]
        return []

    input_ids = [s['id'] for s in sentences]
    input_ids_set = set(input_ids)

    valid_sessions = []
    for index, session in enumerate(sessions):
        raw_indices = session.get('sentence_indices') or session.get('indices') or []
        if not isinstance(raw_indices, list):
            raw_indices = []

        # Check if the model returned 0-based offsets instead of actual IDs
        matching_actual = sum(1 for idx in raw_indices if idx in input_ids_set)
        matching_offsets = sum(1 for idx in raw_indices if 0 <= idx < len(sentences))

        mapped_ids = []
        if matching_offsets > matching_actual:
            # Map 0-based offset to actual ID
            for idx in raw_indices:
                if 0 <= idx < len(sentences):
                    mapped_ids.append(input_ids[idx])
        else:
            # Filter to keep only indices that exist in input_ids
            for idx in raw_indices:
                if idx in input_ids_set:
                    mapped_ids.append(idx)

        # Only keep sessions that have at least one mapped ID
        if mapped_ids:
            mapped_ids.sort()
            valid_sessions.append({
                'title': session.get('title') or f'Session {index + 1}',
                'summary': session.get('summary') or '',
                'mapped_ids': mapped_ids
            })

    if not valid_sessions:
        return [{
            'title': 'Session 1',
            'summary': 'Discussion.',
            'sentence_indices': input_ids
        }]

    # Sort sessions by their first mapped ID
    valid_sessions.sort(key=lambda s: s['mapped_ids'][0])

    # Deduplicate sessions by their start ID
    unique_sessions = []
    seen_start_ids = set()
    for session in valid_sessions:
        start_id = session['mapped_ids'][0]
        if start_id not in seen_start_ids:
            seen_start_ids.add(start_id)
            unique_sessions.append(session)

    # Re-assign every input ID to exactly one session
    final_sessions = [{
        'title': session['title'],
        'summary': session['summary'],
        'sentence_indices': []
    } for session in unique_sessions]

    for id in input_ids:
        assigned_session_idx = 0
        for j in range(1, len(unique_sessions)):
            if unique_sessions[j]['mapped_ids'][0] <= id:
                assigned_session_idx = j
            else:
                break
        final_sessions[assigned_session_idx]['sentence_indices'].append(id)

    # Filter out any sessions that might have ended up with 0 sentences
    return [s for s in final_sessions if len(s['sentence_indices']) > 0]


@app.route('/api/validate-key', methods=['POST'])
def validate_key():
    data = request.json or {}
    api_key = data.get('api_key')
    is_override = True
    
    if not api_key:
        api_key = os.getenv('OPENROUTER_API_KEY') or os.getenv('GEMINI_API_KEY')
        is_override = False
        
    if not api_key:
        return jsonify({
            'valid': False,
            'has_key': False,
            'is_override': is_override,
            'error': 'No API key configured.'
        }), 200
        
    import urllib.request
    import urllib.error
    import json
    
    url = "https://openrouter.ai/api/v1/key"
    headers = {
        'Authorization': f'Bearer {api_key}',
        'HTTP-Referer': 'https://github.com/AnshulUpgrad/PPS-Timestamp-AI-Tool-Upgrad-V2',
        'X-Title': 'PPS Timestamp AI Tool'
    }
    
    req = urllib.request.Request(url, headers=headers, method='GET')
    try:
        with urllib.request.urlopen(req, timeout=10) as response:
            res_data = json.loads(response.read().decode('utf-8'))
            key_data = res_data.get('data', {})

            # Management keys are valid OpenRouter credentials but cannot call
            # completion endpoints, so reject them with a useful explanation.
            if key_data.get('is_management_key') or key_data.get('is_provisioning_key'):
                return jsonify({
                    'valid': False,
                    'has_key': True,
                    'is_override': is_override,
                    'error': 'This is a management/provisioning key. Use a standard OpenRouter API key for AI generation.',
                    'data': key_data
                }), 200

            return jsonify({
                'valid': True,
                'has_key': True,
                'is_override': is_override,
                'data': key_data
            }), 200
    except urllib.error.HTTPError as e:
        try:
            error_body = json.loads(e.read().decode('utf-8'))
            error_message = error_body.get('error', {}).get('message') or error_body.get('message')
        except Exception:
            error_message = None

        # Only an authentication response proves the key is invalid. Other
        # statuses can be temporary OpenRouter/rate-limit failures.
        is_auth_failure = e.code == 401
        return jsonify({
            'valid': False if is_auth_failure else None,
            'has_key': True,
            'is_override': is_override,
            'error': error_message or (
                'OpenRouter rejected this API key.' if is_auth_failure
                else f'OpenRouter verification is temporarily unavailable (HTTP {e.code}).'
            )
        }), 200
    except Exception as e:
        return jsonify({
            'valid': None,
            'has_key': True,
            'is_override': is_override,
            'error': f'Could not reach OpenRouter to verify the key: {str(e)}'
        }), 200

@app.route('/api/chunk-sessions', methods=['POST'])
def chunk_sessions():
    data = request.json or {}
    sentences = data.get('sentences', [])
    model_name = normalize_model_name(data.get('model'))
    single_batch = bool(data.get('single_batch'))
        
    # API key from headers, request body, or environment
    api_key = get_openrouter_api_key(data)
    
    if not api_key:
        return jsonify({'error': 'OpenRouter API Key is missing. Please configure OPENROUTER_API_KEY in your .env file.'}), 400
        
    if not sentences:
        return jsonify({'error': 'No sentences provided for chunking.'}), 400
        
    # Construct the iterative batching logic
    BATCH_SIZE = 20
    sessions = []
    
    # Helper to chunk a specific batch of sentences using OpenRouter API
    def call_ai_chunker(batch_sentences):
        sentences_str = ""
        for s in batch_sentences:
            sentences_str += f"[{s['id']}] {s['text']}\n"
            
        first_idx = batch_sentences[0]['id']
        last_idx = batch_sentences[-1]['id']
        
        prompt_tmpl = load_prompt_template('session_chunking.md')
        prompt = prompt_tmpl.replace("{first_index}", str(first_idx)).replace("{last_index}", str(last_idx)).replace("{sentences_str}", sentences_str)
        # Fallback for old templates
        prompt = prompt.replace("from 0 to", f"from {first_idx} to")
        
        schema = {
            "type": "object",
            "properties": {
                "sessions": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "title": {"type": "string"},
                            "summary": {"type": "string"},
                            "sentence_indices": {
                                "type": "array",
                                "items": {"type": "integer"}
                            }
                        },
                        "required": ["title", "summary", "sentence_indices"],
                        "additionalProperties": False
                    }
                }
            },
            "required": ["sessions"],
            "additionalProperties": False
        }
        
        max_retries = 3
        last_error = None
        
        for attempt in range(max_retries):
            try:
                text_response = call_openrouter_api(
                    model_name=model_name,
                    prompt=prompt,
                    response_schema=schema,
                    api_key=api_key
                )
                
                res_data = clean_json_response(text_response)
                
                # Validate response with Pydantic schema
                validated_data = ChunkingResult.model_validate(res_data)
                
                return [session.model_dump() for session in validated_data.sessions]
            except Exception as e:
                print(f"Chunking generation attempt {attempt + 1} failed: {e}")
                last_error = e
                import time
                time.sleep(1)
                
        raise last_error

    if single_batch:
        try:
            try:
                raw_sessions = call_ai_chunker(sentences)
                reconciled = reconcile_chunks(raw_sessions, sentences)
                return jsonify({'sessions': reconciled}), 200
            except FileNotFoundError as tmpl_err:
                return jsonify({'error': f'Prompt template file missing on server: {str(tmpl_err)}. Make sure the markdown_files/ directory is present.'}), 500
        except OpenRouterError as e:
            return jsonify({'error': str(e)}), e.status_code
        except Exception as e:
            return jsonify({'error': f"Failed to run chunking: {str(e)}"}), 500
        
    i = 0
    total_sentences = len(sentences)
    
    while i < total_sentences:
        current_batch_sentences = sentences[i:i + BATCH_SIZE]
        
        if not sessions:
            # First batch: just chunk the first 40 sentences
            batch_to_chunk = current_batch_sentences
            i += BATCH_SIZE
        else:
            # Bundle the last 2 chunks from the previous batch with the next 40 sentences
            last_chunks_count = min(2, len(sessions))
            last_two_chunks = sessions[-last_chunks_count:]
            
            # Extract sentence indices from those last chunks
            bundled_indices = []
            for chunk in last_two_chunks:
                bundled_indices.extend(chunk.get('sentence_indices', []))
                
            # Get actual sentence dicts for those indices
            bundled_sentences = [s for s in sentences if s['id'] in bundled_indices]
            
            # Combine the last 2 chunks' sentences with the next 40 sentences
            batch_to_chunk = bundled_sentences + current_batch_sentences
            
            # Remove the last chunks from sessions since they are being re-chunked
            sessions = sessions[:-last_chunks_count]
            
            # Advance our pointer
            i += BATCH_SIZE
            
        try:
            try:
                new_sessions = call_ai_chunker(batch_to_chunk)
            except FileNotFoundError as tmpl_err:
                return jsonify({'error': f'Prompt template file missing on server: {str(tmpl_err)}. Make sure the markdown_files/ directory is present.'}), 500
            sessions.extend(new_sessions)
        except OpenRouterError as e:
            return jsonify({'error': f"OpenRouter Error at sentence index {i - BATCH_SIZE}: {str(e)}"}), e.status_code
        except Exception as e:
            return jsonify({'error': f"Failed to run chunking at sentence index {i - BATCH_SIZE}: {str(e)}"}), 500
            
    reconciled = reconcile_chunks(sessions, sentences)
    return jsonify({'sessions': reconciled}), 200

@app.route('/api/chunks/<path:filename>', methods=['GET'])
def get_chunks(filename):
    chunks_filename = f"{filename}_chunks.json"
    chunks_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], chunks_filename)
    
    deleted_filename = f"{filename}_deleted_sentences.json"
    deleted_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], deleted_filename)
    
    deleted_sentences = []
    deleted_words = []
    if os.path.exists(deleted_path) and os.path.isfile(deleted_path):
        try:
            with open(deleted_path, 'r', encoding='utf-8') as f:
                del_data = json.load(f)
                deleted_sentences = del_data.get('deleted_sentences', [])
                deleted_words = del_data.get('deleted_words', [])
        except Exception as e:
            print(f"Failed to load deleted sentences: {e}")
    
    if os.path.exists(chunks_path) and os.path.isfile(chunks_path):
        try:
            with open(chunks_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            # Fallback to deleted_sentences inside chunks JSON if separate file doesn't exist
            if not deleted_sentences:
                deleted_sentences = data.get('deleted_sentences', [])
            if not deleted_words:
                deleted_words = data.get('deleted_words', [])
            return jsonify({
                'exists': True,
                'sessions': data.get('sessions', []),
                'deleted_sentences': deleted_sentences,
                'deleted_words': deleted_words,
                'updated_at': data.get('updated_at', '')
            }), 200
        except Exception as e:
            return jsonify({
                'exists': False,
                'error': f"Failed to load chunks: {str(e)}"
            }), 500
            
    return jsonify({'exists': False}), 200

@app.route('/api/save-chunks/<path:filename>', methods=['POST'])
def save_chunks(filename):
    data = request.json or {}
    sessions = data.get('sessions', [])
    deleted_sentences = data.get('deleted_sentences', [])
    deleted_words = data.get('deleted_words', [])

    if data.get('confirm'):
        empty_sessions = [
            index + 1 for index, session in enumerate(sessions)
            if not session.get('sentence_indices')
        ]
        if empty_sessions:
            return jsonify({
                'error': f"Cannot confirm empty session(s): {', '.join(map(str, empty_sessions))}."
            }), 400
    
    chunks_filename = f"{filename}_chunks.json"
    chunks_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], chunks_filename)
    
    deleted_filename = f"{filename}_deleted_sentences.json"
    deleted_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], deleted_filename)
    
    try:
        # Check if raw transcript was sent directly (e.g. from serverless Modal flow)
        raw_transcript = data.get('raw_transcript')
        if raw_transcript:
            transcript_filename = f"{filename}.json"
            transcript_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], transcript_filename)
            with open(transcript_path, 'w', encoding='utf-8') as f:
                json.dump(raw_transcript, f, ensure_ascii=False, indent=2)

        import datetime
        save_data = {
            'filename': filename,
            'sessions': sessions,
            'deleted_sentences': deleted_sentences,
            'deleted_words': deleted_words,
            'updated_at': datetime.datetime.now().isoformat()
        }
        with open(chunks_path, 'w', encoding='utf-8') as f:
            json.dump(save_data, f, ensure_ascii=False, indent=2)
            
        # Save deleted sentences to a separate dedicated JSON file
        deleted_data = {
            'filename': filename,
            'deleted_sentences': deleted_sentences,
            'deleted_words': deleted_words,
            'updated_at': datetime.datetime.now().isoformat()
        }
        with open(deleted_path, 'w', encoding='utf-8') as f:
            json.dump(deleted_data, f, ensure_ascii=False, indent=2)
            
        return jsonify({
            'message': 'Chunks and transcript deletions saved successfully',
            'filename': chunks_filename
        }), 200
    except Exception as e:
        return jsonify({'error': f"Failed to save chunks: {str(e)}"}), 500

def parse_time_increment(val):
    if not val or not isinstance(val, str) or not val.strip():
        return 0.0
    
    val = val.strip()
    is_negative = False
    if val.startswith('-'):
        is_negative = True
        val = val[1:]
    elif val.startswith('+'):
        val = val[1:]
        
    parts = val.split(':')
    try:
        if len(parts) == 3:
            h = int(parts[0])
            m = int(parts[1])
            s = float(parts[2])
            seconds = h * 3600.0 + m * 60.0 + s
        elif len(parts) == 2:
            m = int(parts[0])
            s = float(parts[1])
            seconds = m * 60.0 + s
        elif len(parts) == 1:
            seconds = float(parts[0])
        else:
            seconds = 0.0
            
        return -seconds if is_negative else seconds
    except ValueError:
        return 0.0

@app.route('/api/export-docx/<path:filename>', methods=['POST'])
def export_docx(filename):
    try:
        import io
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from flask import send_file
    except ImportError as e:
        return jsonify({'error': f'docx library not installed or import failed: {str(e)}'}), 500
        
    data = request.json or {}
    sessions_data = data.get('sessions', [])
    
    # Pull deleted sentences directly from the separate JSON file on disk
    deleted_sentences = []
    deleted_words = []
    deleted_filename = f"{filename}_deleted_sentences.json"
    deleted_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], deleted_filename)
    if os.path.exists(deleted_path) and os.path.isfile(deleted_path):
        try:
            with open(deleted_path, 'r', encoding='utf-8') as f:
                del_data = json.load(f)
                deleted_sentences = del_data.get('deleted_sentences', [])
                deleted_words = del_data.get('deleted_words', [])
        except Exception as e:
            print(f"Error loading deleted sentences from disk during DOCX export: {e}")
            
    # Fallback to payload if file was not found or failed to load
    if not deleted_sentences:
        deleted_sentences = data.get('deleted_sentences', [])
    if not deleted_words:
        deleted_words = data.get('deleted_words', [])
    
    if not sessions_data:
        return jsonify({'error': 'No session data provided for export.'}), 400
        
    # Attempt to load original sentences to construct the transcript text
    t_data = data.get('raw_transcript')
    if not t_data:
        transcript_path = os.path.join(app.config['TRANSCRIPTIONS_FOLDER'], f"{filename}.json")
        if os.path.exists(transcript_path):
            try:
                with open(transcript_path, 'r', encoding='utf-8') as f:
                    t_data = json.load(f)
            except Exception as e:
                print(f"Error loading transcript for DOCX from disk: {e}")
                
    sentences_map = {}
    if t_data:
        try:
            sentences_list = split_transcript_into_sentences(t_data)
            sentences_map = {s['id']: s for s in sentences_list}
        except Exception as e:
            print(f"Error parsing transcript for DOCX: {e}")

    # Apply word-level deletions to the reconstructed transcript before export.
    for deleted_word in deleted_words:
        sentence = sentences_map.get(deleted_word.get('sentence_id'))
        if not sentence or not sentence.get('words'):
            continue
        target_start = float(deleted_word.get('start', -1.0))
        target_end = float(deleted_word.get('end', -1.0))
        target_text = str(deleted_word.get('word', '')).strip()
        removed = False
        kept_words = []
        for word in sentence.get('words', []):
            same_time = (
                abs(float(word.get('start', -2.0)) - target_start) < 0.002
                and abs(float(word.get('end', -2.0)) - target_end) < 0.002
            )
            same_text = str(word.get('word', '')).strip() == target_text
            if not removed and same_time and same_text:
                removed = True
                continue
            kept_words.append(word)
        sentence['words'] = kept_words
        text = " ".join(str(word.get('word', '')).strip() for word in kept_words if str(word.get('word', '')).strip())
        text = re.sub(r'\s+([,.;:!?%])', r'\1', text).strip()
        sentence['text'] = text

    # Estimate video duration in seconds
    duration_sec = 0.0
    if sentences_map:
        for s in sentences_map.values():
            duration_sec = max(duration_sec, s.get('end', 0.0))
            
    if duration_sec == 0.0 and sessions_data:
        max_session_ts = 0.0
        for s in sessions_data:
            max_session_ts = max(max_session_ts, abs(s.get('end', 0.0)))
        if max_session_ts > 10000.0:
            duration_sec = max_session_ts / 1000.0
        else:
            duration_sec = max_session_ts
            
    if duration_sec == 0.0:
        duration_sec = 3600.0  # Fallback to 1 hour
        
    # Detect and normalize deleted_sentences if they are in milliseconds
    max_deleted_ts = 0.0
    if deleted_sentences:
        for ds in deleted_sentences:
            max_deleted_ts = max(max_deleted_ts, abs(ds.get('start', 0.0)), abs(ds.get('end', 0.0)))
    if max_deleted_ts > (duration_sec + 10.0):
        for ds in deleted_sentences:
            if 'start' in ds:
                ds['start'] = ds.get('start', 0.0) / 1000.0
            if 'end' in ds:
                ds['end'] = ds.get('end', 0.0) / 1000.0

    # Detect and normalize sessions start/end if they are in milliseconds
    max_session_ts = 0.0
    if sessions_data:
        for s in sessions_data:
            max_session_ts = max(max_session_ts, abs(s.get('start', 0.0)), abs(s.get('end', 0.0)))
    if max_session_ts > (duration_sec + 10.0):
        for s in sessions_data:
            if 'start' in s:
                s['start'] = s.get('start', 0.0) / 1000.0
            if 'end' in s:
                s['end'] = s.get('end', 0.0) / 1000.0

    # Detect and normalize visuals items and details if they are in milliseconds
    max_visual_ts = 0.0
    if sessions_data:
        for s in sessions_data:
            visuals = s.get('visuals', {})
            if visuals:
                v_content = visuals.get('content', {})
                if v_content:
                    for item in v_content.get('items', []):
                        max_visual_ts = max(max_visual_ts, abs(item.get('timestamp', 0.0)))
                    for detail in v_content.get('details', []):
                        max_visual_ts = max(max_visual_ts, abs(detail.get('timestamp', 0.0)))
    if max_visual_ts > (duration_sec + 10.0):
        for s in sessions_data:
            visuals = s.get('visuals', {})
            if visuals:
                v_content = visuals.get('content', {})
                if v_content:
                    v_items = v_content.get('items', [])
                    for item in v_items:
                        if 'timestamp' in item:
                            item['timestamp'] = item.get('timestamp', 0.0) / 1000.0
                    v_details = v_content.get('details', [])
                    for detail in v_details:
                        if 'timestamp' in detail:
                            detail['timestamp'] = detail.get('timestamp', 0.0) / 1000.0

    # Parse time increment if provided and shift timestamps
    time_increment = data.get('time_increment', '')
    time_shift = parse_time_increment(time_increment)

    if time_shift != 0:
        # 1. Shift deleted_sentences
        if deleted_sentences:
            for ds in deleted_sentences:
                if 'start' in ds:
                    ds['start'] = ds.get('start', 0.0) + time_shift
                if 'end' in ds:
                    ds['end'] = ds.get('end', 0.0) + time_shift

        if deleted_words:
            for word in deleted_words:
                if 'start' in word:
                    word['start'] = word.get('start', 0.0) + time_shift
                if 'end' in word:
                    word['end'] = word.get('end', 0.0) + time_shift

        # 2. Shift sentences_map
        if sentences_map:
            for s in sentences_map.values():
                if 'start' in s:
                    s['start'] = s.get('start', 0.0) + time_shift
                if 'end' in s:
                    s['end'] = s.get('end', 0.0) + time_shift

        # 3. Shift sessions_data start/end, and slide/visual content timestamps
        if sessions_data:
            for s in sessions_data:
                if 'start' in s:
                    s['start'] = s.get('start', 0.0) + time_shift
                if 'end' in s:
                    s['end'] = s.get('end', 0.0) + time_shift
                    
                visuals = s.get('visuals', {})
                if visuals:
                    v_content = visuals.get('content', {})
                    if v_content:
                        if 'heading_timestamp' in v_content:
                            v_content['heading_timestamp'] = v_content.get('heading_timestamp', 0.0) + time_shift
                        v_items = v_content.get('items', [])
                        for item in v_items:
                            if 'timestamp' in item:
                                item['timestamp'] = item.get('timestamp', 0.0) + time_shift
                        v_details = v_content.get('details', [])
                        for detail in v_details:
                            if 'timestamp' in detail:
                                detail['timestamp'] = detail.get('timestamp', 0.0) + time_shift

    # Helper functions for styling
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
        # top, bottom, left, right are in dxa (1 pt = 20 dxa)
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_cell_borders(cell, color="E2E8F0", sz="4", val="single"):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), val)
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def format_seconds(seconds):
        if seconds is None:
            return "0:00"
        val = float(seconds)
        
        # Round to the nearest second
        val = int(val + 0.5) if val >= 0 else int(val - 0.5)
        
        is_neg = val < 0
        val = abs(val)
        hours = int(val // 3600)
        mins = int((val % 3600) // 60)
        secs = int(val % 60)
        prefix = "-" if is_neg else ""
        if hours > 0:
            return f"{prefix}{hours:02d}:{mins:02d}:{secs:02d}"
        else:
            return f"{prefix}{mins}:{secs:02d}"

    # Create document
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Key Points & Visual Layout Export")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"Source file: {filename}")
    subtitle_run.font.name = 'Calibri'
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacer
    
    # Create Table
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Column width setting (in inches)
    col_widths = [Inches(2.6), Inches(3.9)]
    
    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = ''
    
    # Configure Header Cell 1
    p1 = hdr_cells[0].paragraphs[0]
    p1_run = p1.add_run("Plate & Transcript Content")
    p1_run.font.name = 'Calibri'
    p1_run.font.size = Pt(11)
    p1_run.font.bold = True
    p1_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Configure Header Cell 2
    p2 = hdr_cells[1].paragraphs[0]
    p2_run = p2.add_run("Curriculum Highlights & Visual Mapping")
    p2_run.font.name = 'Calibri'
    p2_run.font.size = Pt(11)
    p2_run.font.bold = True
    p2_run.font.color.rgb = RGBColor(255, 255, 255)
    
    for i, cell in enumerate(hdr_cells):
        set_cell_background(cell, "1E293B") # Slate 800
        set_cell_margins(cell, top=180, bottom=180, left=180, right=180)
        set_cell_borders(cell, color="475569") # Slate 600
        cell.width = col_widths[i]
        
    # Reconstruct original indices for each session to find original boundaries (including deleted sentences)
    session_original_ranges = []
    session_original_indices = []
    for s in sessions_data:
        session_original_indices.append(list(s.get('sentence_indices', [])))
        
    for ds in deleted_sentences:
        ds_id = ds.get('id')
        if ds_id is None:
            continue
        try:
            ds_id = int(ds_id)
        except ValueError:
            continue
            
        inserted = False
        # 1. Check if it fits inside the active range of any session
        for indices in session_original_indices:
            if indices:
                min_id = min(indices)
                max_id = max(indices)
                if min_id <= ds_id <= max_id:
                    indices.append(ds_id)
                    inserted = True
                    break
                    
        # 2. Check if it is adjacent to any session
        if not inserted:
            for indices in session_original_indices:
                if indices:
                    min_id = min(indices)
                    max_id = max(indices)
                    if ds_id == max_id + 1 or ds_id == min_id - 1:
                        indices.append(ds_id)
                        inserted = True
                        break
                        
        # 3. Fallback to last session
        if not inserted and session_original_indices:
            session_original_indices[-1].append(ds_id)
            
    # Calculate original start and end times for each session
    for indices in session_original_indices:
        indices.sort()
        orig_start = 0.0
        orig_end = 0.0
        if indices and sentences_map:
            first_sent = sentences_map.get(indices[0])
            last_sent = sentences_map.get(indices[-1])
            if first_sent:
                orig_start = first_sent.get('start', 0.0)
            if last_sent:
                orig_end = last_sent.get('end', 0.0)
        session_original_ranges.append((orig_start, orig_end))

    # Populate Rows
    for idx, session in enumerate(sessions_data):
        row = table.add_row()
        row_cells = row.cells
        
        # Widths
        row_cells[0].width = col_widths[0]
        row_cells[1].width = col_widths[1]
        
        # Margins & borders
        for cell in row_cells:
            set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
            set_cell_borders(cell, color="CBD5E1") # Slate 300 (light gray)
            
        # Left Cell Shading (very light blue/gray for Session Column)
        set_cell_background(row_cells[0], "F8FAFC") # Slate 50
        
        # Time duration
        session_start = 0.0
        session_end = 0.0
        indices = session.get('sentence_indices', [])
        if indices and sentences_map:
            first_sent = sentences_map.get(indices[0])
            last_sent = sentences_map.get(indices[-1])
            if first_sent:
                session_start = first_sent.get('start', 0.0)
            if last_sent:
                session_end = last_sent.get('end', 0.0)
        duration_str = f"Time: {format_seconds(session_start)} - {format_seconds(session_end)}"
        
        # Visual Template
        visuals = session.get('visuals', {})
        template_name = visuals.get('template_name', 'Face Only')
        why_chosen = visuals.get('why_chosen', '')
        graphics_req = visuals.get('graphics_required', False)
        v_content = visuals.get('content', {})
        v_title = v_content.get('title', '')
        v_items = v_content.get('items', [])
        v_details = v_content.get('details', [])
        heading_timestamp = v_content.get('heading_timestamp', session_start)
        heading_timestamp = max(session_start, min(session_end, float(heading_timestamp)))
        heading_timestamp_str = format_seconds(heading_timestamp)

        # Check for deleted sentences falling within this session's original range
        session_deleted = []
        orig_start, orig_end = session_original_ranges[idx]
        for ds in deleted_sentences:
            ds_start = ds.get('start', 0.0)
            ds_end = ds.get('end', 0.0)
            if orig_start - 0.15 <= ds_start <= orig_end + 0.15:
                session_deleted.append(f"{format_seconds(ds_start)} - {format_seconds(ds_end)}")
        for word in deleted_words:
            word_start = word.get('start', 0.0)
            word_end = word.get('end', 0.0)
            if orig_start - 0.15 <= word_start <= orig_end + 0.15:
                session_deleted.append(
                    f"{format_seconds(word_start)} - {format_seconds(word_end)} ({word.get('word', '').strip()})"
                )

        # Left Cell Content
        cell_left = row_cells[0]
        p_session = cell_left.paragraphs[0]
        p_session.paragraph_format.space_after = Pt(4)
        run_s_title = p_session.add_run(f"Plate {idx + 1}")
        run_s_title.font.name = 'Calibri'
        run_s_title.font.size = Pt(11)
        run_s_title.font.bold = True
        run_s_title.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        
        # Plain text timestamp beside heading
        run_s_ts = p_session.add_run(f" {heading_timestamp_str}")
        run_s_ts.font.name = 'Calibri'
        run_s_ts.font.size = Pt(11)
        run_s_ts.font.bold = False
        run_s_ts.font.color.rgb = RGBColor(30, 41, 59)
        
        if session_deleted:
            run_deleted = p_session.add_run(f" [{', '.join(session_deleted)}]")
            run_deleted.font.name = 'Calibri'
            run_deleted.font.size = Pt(10)
            run_deleted.font.bold = True
            run_deleted.font.color.rgb = RGBColor(220, 38, 38) # Red 600
        
        p_time = cell_left.add_paragraph()
        p_time.paragraph_format.space_after = Pt(8)
        run_time = p_time.add_run(duration_str)
        run_time.font.name = 'Calibri'
        run_time.font.size = Pt(9.5)
        run_time.font.italic = True
        run_time.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
        
        # Transcript Content
        session_sentences = []
        for sid in indices:
            if sid in sentences_map:
                session_sentences.append(sentences_map[sid].get('text', '').strip())
        session_text = " ".join(session_sentences)
        
        p_trans_hdr = cell_left.add_paragraph()
        p_trans_hdr.paragraph_format.space_after = Pt(2)
        run_trans_hdr = p_trans_hdr.add_run("Transcript:")
        run_trans_hdr.font.name = 'Calibri'
        run_trans_hdr.font.size = Pt(9)
        run_trans_hdr.font.bold = True
        run_trans_hdr.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
        
        p_trans_body = cell_left.add_paragraph()
        p_trans_body.paragraph_format.line_spacing = 1.15
        p_trans_body.paragraph_format.space_after = Pt(0)
        run_trans_body = p_trans_body.add_run(session_text or "No transcript text available.")
        run_trans_body.font.name = 'Calibri'
        run_trans_body.font.size = Pt(9.5)
        run_trans_body.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
        
        # Right Cell Content
        cell_right = row_cells[1]
        
        p_temp = cell_right.paragraphs[0]
        p_temp.paragraph_format.space_before = Pt(0)
        p_temp.paragraph_format.space_after = Pt(2)
        run_temp_lbl = p_temp.add_run("Visual Layout Suggestion: ")
        run_temp_lbl.font.name = 'Calibri'
        run_temp_lbl.font.size = Pt(9.5)
        run_temp_lbl.font.bold = True
        run_temp_lbl.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
        
        run_temp_val = p_temp.add_run(template_name)
        run_temp_val.font.name = 'Calibri'
        run_temp_val.font.size = Pt(10)
        run_temp_val.font.bold = True
        run_temp_val.font.color.rgb = RGBColor(13, 148, 136) # Teal 600
        
        if graphics_req:
            run_req = p_temp.add_run(" (Graphics Required)")
            run_req.font.name = 'Calibri'
            run_req.font.size = Pt(9)
            run_req.font.italic = True
            run_req.font.color.rgb = RGBColor(220, 38, 38) # Red 600
            
            
        # Visual Content details (only if template is not Face Only)
        if template_name == "Name aston":
            aston_name = visuals.get('aston_name', '')
            p_aston = cell_right.add_paragraph()
            p_aston.paragraph_format.space_before = Pt(6)
            p_aston.paragraph_format.space_after = Pt(2)
            run_aston_lbl = p_aston.add_run("Aston Name: ")
            run_aston_lbl.font.name = 'Calibri'
            run_aston_lbl.font.size = Pt(9.5)
            run_aston_lbl.font.bold = True
            run_aston_lbl.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
            
            run_aston_val = p_aston.add_run(aston_name)
            run_aston_val.font.name = 'Calibri'
            run_aston_val.font.size = Pt(10)
            run_aston_val.font.bold = True
            run_aston_val.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        elif template_name == "Custom Template":
            custom_text = visuals.get('custom_text', '')
            p_custom = cell_right.add_paragraph()
            p_custom.paragraph_format.space_before = Pt(6)
            p_custom.paragraph_format.space_after = Pt(2)
            run_custom_lbl = p_custom.add_run("Custom Text: ")
            run_custom_lbl.font.name = 'Calibri'
            run_custom_lbl.font.size = Pt(9.5)
            run_custom_lbl.font.bold = True
            run_custom_lbl.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
            
            run_custom_val = p_custom.add_run(custom_text)
            run_custom_val.font.name = 'Calibri'
            run_custom_val.font.size = Pt(10)
            run_custom_val.font.bold = True
            run_custom_val.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        elif template_name != "Face Only" and (v_title or v_items or v_details):
            p_vc_hdr = cell_right.add_paragraph()
            p_vc_hdr.paragraph_format.space_before = Pt(6)
            p_vc_hdr.paragraph_format.space_after = Pt(2)
            run_vc_hdr = p_vc_hdr.add_run("Visual Content Details:")
            run_vc_hdr.font.name = 'Calibri'
            run_vc_hdr.font.size = Pt(9.5)
            run_vc_hdr.font.bold = True
            run_vc_hdr.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
            
            if v_title:
                p_v_title = cell_right.add_paragraph()
                p_v_title.paragraph_format.space_after = Pt(3)
                run_v_title_lbl = p_v_title.add_run("  Heading: ")
                run_v_title_lbl.font.name = 'Calibri'
                run_v_title_lbl.font.size = Pt(9)
                run_v_title_lbl.font.bold = False
                run_v_title_lbl.font.color.rgb = RGBColor(0, 0, 0)
                
                run_v_title_val = p_v_title.add_run(v_title)
                run_v_title_val.font.name = 'Calibri'
                run_v_title_val.font.size = Pt(9.5)
                run_v_title_val.font.bold = True
                run_v_title_val.font.color.rgb = RGBColor(0, 0, 0)
                
                run_v_title_ts = p_v_title.add_run(f" {heading_timestamp_str}")
                run_v_title_ts.font.name = 'Calibri'
                run_v_title_ts.font.size = Pt(9.5)
                run_v_title_ts.font.bold = False
                run_v_title_ts.font.color.rgb = RGBColor(0, 0, 0)
                
            # Render all pointer rows chronologically, even when a template uses
            # both list items and labelled details.
            visual_entries = [
                ('item', item) for item in v_items
            ] + [
                ('detail', detail) for detail in v_details
            ]
            visual_entries.sort(key=lambda entry: float(entry[1].get('timestamp', session_start)))
            is_differentiation = template_name.startswith('Differentiation Template')

            for entry_type, entry in visual_entries:
                paragraph_style = None if is_differentiation else 'List Bullet 2'
                p_entry = cell_right.add_paragraph(style=paragraph_style)
                p_entry.paragraph_format.space_after = Pt(2)
                entry_timestamp = entry.get('timestamp', session_start)

                run_ts = p_entry.add_run(f"[{format_seconds(entry_timestamp)}] ")
                run_ts.font.name = 'Calibri'
                run_ts.font.size = Pt(8.5)
                run_ts.font.bold = False
                run_ts.font.color.rgb = RGBColor(0, 0, 0)

                if entry_type == 'detail':
                    run_lbl = p_entry.add_run(f"{entry.get('label', '')}. " if is_differentiation else f"{entry.get('label', '')}: ")
                    run_lbl.font.name = 'Calibri'
                    run_lbl.font.size = Pt(9)
                    run_lbl.font.bold = False if is_differentiation else True
                    run_lbl.font.color.rgb = RGBColor(0, 0, 0)
                    entry_value = entry.get('value', '')
                else:
                    entry_value = entry.get('value', '')

                run_val = p_entry.add_run(entry_value)
                run_val.font.name = 'Calibri'
                run_val.font.size = Pt(9)
                run_val.font.bold = True
                run_val.font.color.rgb = RGBColor(0, 0, 0)
                    
        # Additional Text Content is omitted from the DOCX file per user request
            
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Return as download attachment
    base_name, _ = os.path.splitext(os.path.basename(filename))
    export_filename = f"{base_name}_curriculum_export.docx"
    
    return send_file(
        file_stream,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=export_filename
    )

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    # Disable reloader and debug mode to prevent background process crashes in Google Colab
    app.run(debug=False, host='0.0.0.0', port=port)
